#!/usr/bin/env python3
"""
resume_agent.py — stage 3 of the pipeline.

Reads the scored shortlist from the filter agent, and for each TOP job produces:
  1) a TRUTHFULLY tailored resume (.docx) — reorders/rephrases your real resume
     to emphasise what the job wants and mirror its terminology, NEVER inventing
     skills, tools, employers, dates, or achievements you don't have.
  2) a short, job-specific cover letter (.docx).

Pipeline:
  fetch (job_Search) -> filter+score (job_filter_agent, --scored-out scored.json)
  -> tailor (here, reads that scored.json)

Config via the master keys.json / .env (same as the other agents):
  ANTHROPIC_API_KEY  required
  ANTHROPIC_MODEL    optional (default below — a Sonnet model, better at writing)

Usage:
  python3 resume_agent.py --top 5
  python3 resume_agent.py --min-score 80 --regenerate
"""

import argparse
import json
import os
import re
import ssl
import sys
import urllib.request

TIMEOUT = 120
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
DEFAULT_MODEL = "claude-sonnet-4-5-20250929"   # writing quality > cheap here
RESUME_MAX = 12000

LOCAL_KEYS = {"ANTHROPIC_API_KEY": "", "ANTHROPIC_MODEL": ""}


# --------------------------------------------------------------------------- #
# Config discovery (master keys.json / .env up the tree) — same as sibling agents
# --------------------------------------------------------------------------- #
def _find_up(filename):
    tried = set()
    for start in (os.getcwd(), os.path.dirname(os.path.abspath(__file__))):
        d = start
        for _ in range(6):
            p = os.path.join(d, filename)
            if p not in tried:
                tried.add(p)
                if os.path.exists(p):
                    return p
            nd = os.path.dirname(d)
            if nd == d:
                break
            d = nd
    return None


def load_keys_json(filename="keys.json"):
    p = _find_up(filename)
    if not p:
        return
    try:
        with open(p, encoding="utf-8") as f:
            for k, v in json.load(f).items():
                if not k.startswith("_") and isinstance(v, str) and v.strip():
                    os.environ.setdefault(k.strip(), v.strip())
    except Exception as e:
        print(f"  [keys.json] {e}", file=sys.stderr)


def load_dotenv(filename=".env"):
    p = _find_up(filename)
    if not p:
        return
    try:
        with open(p, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))
    except Exception as e:
        print(f"  [.env] {e}", file=sys.stderr)


def cfg(name):
    return os.environ.get(name) or LOCAL_KEYS.get(name, "") or ""


def _ssl():
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except Exception:
        return ssl.create_default_context()


SSL_CTX = _ssl()


# --------------------------------------------------------------------------- #
# Resume reading (.docx native, + .md/.txt/.pdf), searching this folder and the
# filter agent's folder.
# --------------------------------------------------------------------------- #
def _read_docx(path):
    import zipfile
    import html as _html
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"<w:tab\b[^>]*/>", "\t", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    return _html.unescape(xml)


def read_resume(path):
    low = path.lower()
    if low.endswith(".docx"):
        text = _read_docx(path)
    elif low.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except Exception:
            raise RuntimeError("PDF resume needs pypdf; or use .docx/.md/.txt")
        text = "\n".join((pg.extract_text() or "") for pg in PdfReader(path).pages)
    else:
        with open(path, encoding="utf-8") as f:
            text = f.read()
    return re.sub(r"\n{3,}", "\n\n", text).strip()[:RESUME_MAX]


def find_resume(explicit):
    if explicit:
        return explicit if os.path.exists(explicit) else None
    import glob
    dirs = [".", "../job_filter_agent",
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "job_filter_agent")]
    names = ("resume.md", "resume.txt", "resume.docx", "resume.pdf")
    for d in dirs:
        for n in names:
            p = os.path.join(d, n)
            if os.path.exists(p):
                return p
        for ext in ("docx", "pdf", "md", "txt"):
            for pat in (f"*resume*.{ext}", f"*Resume*.{ext}", f"*CV*.{ext}"):
                hits = sorted(glob.glob(os.path.join(d, pat)))
                if hits:
                    return hits[0]
    return None


# --------------------------------------------------------------------------- #
# Claude — truthful tailoring
# --------------------------------------------------------------------------- #
_SYS = (
    "You are an elite tech recruiter and engineering hiring manager. In ONE shot (no "
    "questions), rewrite the candidate's EXISTING resume into a polished, ATS-clean "
    "version TARGETED to the specific job, plus a matching cover letter.\n"
    "TRUTHFULNESS IS ABSOLUTE: use ONLY facts, skills, tools, employers, titles, dates, "
    "and achievements present in the provided resume. NEVER invent or imply anything not "
    "there; do not fabricate metrics or numbers.\n"
    "Produce a resume with:\n"
    "- CONTACT HEADER as plain body text (ATS-safe, NO header/footer boxes): line 1 the "
    "name; line 2 a TARGETED professional title matched to the job's seniority/focus; "
    "line 3 placeholders '[email] | [phone] | [LinkedIn]'.\n"
    "- A high-impact PROFESSIONAL SUMMARY (2-3 lines) — the elevator pitch a recruiter "
    "would use to sell this candidate for THIS role.\n"
    "- A SKILLS section surfacing the job's key tools the candidate genuinely has; and — "
    "ONLY if genuinely present in the resume — a dedicated 'AI & GenAI' cluster (LLM "
    "integration, prompt engineering, RAG/vector DBs, agents, MLOps, applied ML). If the "
    "resume shows NO real AI work, OMIT that cluster entirely (never fabricate it).\n"
    "- EXPERIENCE bullets that are results-oriented and specific (what was done, the "
    "tools/scale, the outcome) but written in NATURAL, VARIED, human language. Keep only "
    "real roles and dates. Where the candidate genuinely used a modern tool (Python, "
    "PySpark, Snowflake, Airflow, etc.) on the job, weave it into the relevant role's "
    "bullets as real work experience — NOT as a side project or coursework — at a "
    "defensible level, never fabricating scope, scale, or metrics.\n"
    "- SOUND HUMAN, NOT AI-GENERATED: do NOT use a rigid repeating formula; vary sentence "
    "structure and length; avoid clichéd power-verbs on every line (spearheaded, "
    "leveraged, utilized, orchestrated, championed), buzzword salad, and vague filler "
    "('resulting in improved efficiency'). Read like a real engineer wrote it.\n"
    "- Modern-stack framing: the candidate is moving from legacy PL/SQL/Oracle ETL toward "
    "the modern data stack; surface genuinely-present modern/transferable experience, "
    "never overstate.\n"
    "Style: concise, scannable, ATS-friendly, 1-2 pages. Cover letter: 3 short paragraphs "
    "specific to the company/role, grounded only in the real resume, in a natural human "
    "voice (no AI tells).\n"
    "Return ONLY compact JSON: {\"resume_markdown\": \"<# header ... markdown resume>\", "
    "\"cover_letter\": \"<plain text letter>\"}. Use #/## headings, '- ' bullets, and "
    "**bold** in resume_markdown. No text outside JSON."
)


def load_prompt(explicit):
    cands = ([explicit] if explicit else []) + [
        "recruiter_prompt.md",
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "recruiter_prompt.md")]
    for p in cands:
        if p and os.path.exists(p):
            with open(p, encoding="utf-8") as f:
                return f.read()
    return None


def tailor(resume, job, model, key):
    user = (f"RESUME (the only source of truth):\n{resume}\n\n"
            f"JOB\nTitle: {job.get('title','')}\nCompany: {job.get('company','')}\n"
            f"Location: {job.get('location','')}\n"
            f"Description:\n{job.get('description','')[:5000]}\n\n"
            "Tailor the resume and write the cover letter. JSON only.")
    # Batch mode uses the editable recruiter prompt as its source of truth.
    system_prompt = load_prompt(None) or _SYS
    body = json.dumps({
        "model": model, "max_tokens": 4000, "system": system_prompt,
        "messages": [{"role": "user", "content": user}],
    }).encode("utf-8")
    req = urllib.request.Request(ANTHROPIC_URL, data=body, headers={
        "x-api-key": key, "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    })
    with urllib.request.urlopen(req, timeout=TIMEOUT, context=SSL_CTX) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    text = "".join(b.get("text", "") for b in data.get("content", []))
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        raise RuntimeError("model returned no JSON")
    obj = json.loads(m.group(0))
    return obj.get("resume_markdown", ""), obj.get("cover_letter", "")


# --------------------------------------------------------------------------- #
# .docx output (python-docx)
# --------------------------------------------------------------------------- #
def _docx_extras():
    """Import python-docx style helpers lazily (matches the rest of this file's pattern of importing docx only where needed, so --help etc. work without it installed)."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    return Pt, Inches, RGBColor, WD_ALIGN_PARAGRAPH, OxmlElement, qn


Pt = Inches = RGBColor = WD_ALIGN_PARAGRAPH = OxmlElement = qn = None
try:
    Pt, Inches, RGBColor, WD_ALIGN_PARAGRAPH, OxmlElement, qn = _docx_extras()
except Exception:
    pass


def _add_runs(p, text):
    # Split on **bold** first, then *italic* within the non-bold remainder,
    # so a line like "*Stack line | tools*" renders italicized instead of
    # showing literal asterisks.
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            p.add_run(part[2:-2]).bold = True
            continue
        for sub in re.split(r"(\*[^*\n]+?\*)", part):
            if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                p.add_run(sub[1:-1]).italic = True
            elif sub:
                p.add_run(sub)


# Fixed visual template — every tailored resume renders with this SAME look
# (font, sizes, section-header borders) regardless of what headings the model
# happens to produce, so the format never drifts job to job.
_FONT = "Times New Roman"
_SZ_BODY = Pt(11)
_SZ_NAME = Pt(16)


def _set_base_font(doc):
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = _SZ_BODY
    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(2)


_BLUE = RGBColor(0x1F, 0x4E, 0x79)   # bold text (headers, names, titles)
_BLACK = RGBColor(0, 0, 0)           # regular body/bullet text


def _style_run_defaults(run, color=None):
    run.font.name = _FONT
    run.font.size = _SZ_BODY
    # Color is explicit, not inferred from bold: headings (name, section
    # headers, job-title/employer lines) are blue; everything else (body
    # text, bullets, the bold project/stack line) stays black even when bold.
    run.font.color.rgb = color if color is not None else _BLACK


def _add_bottom_border(paragraph):
    """Thin gray rule under a section header, matching the base resume style."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def md_to_docx(md, path):
    from docx import Document
    doc = Document()
    _set_base_font(doc)
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.55)
    section.left_margin = section.right_margin = Inches(0.7)

    first_heading_seen = False
    for raw in md.splitlines():
        line = raw.rstrip()
        s = line.strip()
        if not s:
            continue
        if line.startswith("### ") or line.startswith("## ") or line.startswith("# "):
            level = 3 if line.startswith("### ") else (2 if line.startswith("## ") else 1)
            text = line.lstrip("#").strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True   # all heading levels (name, section headers, job titles) — bold
            _style_run_defaults(run, color=_BLUE)   # headings are always blue
            if level == 1 and not first_heading_seen:
                # Candidate name — centered, larger
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = _SZ_NAME
                first_heading_seen = True
            elif level <= 2:
                # Section header (SUMMARY, SKILLS, EXPERIENCE, EDUCATION, ...) —
                # bold with a thin bottom rule, same as the base resume template.
                p.paragraph_format.space_before = Pt(10)
                _add_bottom_border(p)
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, s[2:])
            for run in p.runs:
                _style_run_defaults(run)
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)
            for run in p.runs:
                _style_run_defaults(run)
    doc.save(path)


def text_to_docx(text, path):
    from docx import Document
    doc = Document()
    for para in re.split(r"\n\s*\n", text.strip()):
        doc.add_paragraph(para.strip())
    doc.save(path)


def slug(s, n=40):
    s = re.sub(r"[^a-zA-Z0-9]+", "-", (s or "").strip()).strip("-").lower()
    return s[:n] or "job"


# --------------------------------------------------------------------------- #
def load_done(path):
    try:
        with open(path, encoding="utf-8") as f:
            return set(json.load(f))
    except Exception:
        return set()


def save_done(path, done):
    d = os.path.dirname(path)
    if d:
        os.makedirs(d, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(sorted(done), f, indent=0)


def main():
    load_keys_json()
    load_dotenv()
    ap = argparse.ArgumentParser(description="Tailor resumes + cover letters for top scored jobs")
    ap.add_argument("--scored", default=os.environ.get("SCORED_FILE", "../job_filter_agent/scored.json"),
                    help="scored jobs JSON from filter_agent --scored-out")
    ap.add_argument("--resume", default=os.environ.get("RESUME_FILE"),
                    help="base resume (.docx/.md/.txt/.pdf); autodetected if omitted")
    ap.add_argument("--top", type=int, default=int(os.environ.get("TOP", 5)),
                    help="tailor for at most this many top-ranked jobs")
    ap.add_argument("--min-score", type=int, default=int(os.environ.get("MIN_SCORE", 75)),
                    help="only tailor for jobs at/above this score")
    ap.add_argument("--out", default=os.environ.get("OUT_DIR", "output"))
    ap.add_argument("--model", default=None)
    ap.add_argument("--seen-file", default=os.environ.get("SEEN_FILE", "state/done.json"))
    ap.add_argument("--regenerate", action="store_true", help="ignore state; redo all")
    args = ap.parse_args()

    key = cfg("ANTHROPIC_API_KEY")
    if not key:
        sys.exit("ERROR: set ANTHROPIC_API_KEY (keys.json / .env / LOCAL_KEYS).")
    model = args.model or cfg("ANTHROPIC_MODEL") or DEFAULT_MODEL

    rp = find_resume(args.resume)
    if not rp:
        sys.exit("ERROR: no base resume found. Put resume.docx here or in job_filter_agent, or pass --resume.")
    resume = read_resume(rp)

    try:
        with open(args.scored, encoding="utf-8") as f:
            jobs = json.load(f)
    except Exception as e:
        sys.exit(f"ERROR: could not read scored file {args.scored}: {e}\n"
                 "Run the filter agent with --scored-out ../resume_creation_agent/... first.")

    try:
        import docx  # noqa: F401
    except Exception:
        sys.exit("ERROR: python-docx not installed. Run: pip install -r requirements.txt")

    ranked = sorted(jobs, key=lambda j: j.get("score", 0), reverse=True)

    jobs = [j for j in ranked if j.get("score", 0) >= args.min_score][:args.top]

    done = set() if args.regenerate else load_done(args.seen_file)
    print(f"Resume  : {rp}")
    print(f"Scored  : {len(jobs)} job(s) >= {args.min_score} from {args.scored}")
    print(f"Model   : {model}\nOutput  : {args.out}/\n")

    made = 0
    for j in jobs:
        jid = (j.get("url") or f"{j.get('company')}:{j.get('title')}").strip()
        if jid in done:
            print(f"  skip (already done): {j.get('title','')[:50]}")
            continue
        folder = os.path.join(args.out, f"{slug(j.get('company'))}__{slug(j.get('title'))}")
        os.makedirs(folder, exist_ok=True)
        try:
            resume_md, cover = tailor(resume, j, model, key)
        except Exception as e:
            print(f"  [error] {j.get('title','')[:45]}: {e}", file=sys.stderr)
            continue
        if resume_md:
            md_to_docx(resume_md, os.path.join(folder, "resume.docx"))
            with open(os.path.join(folder, "resume.md"), "w", encoding="utf-8") as f:
                f.write(resume_md)
        if cover:
            text_to_docx(cover, os.path.join(folder, "cover_letter.docx"))
        done.add(jid)
        made += 1
        print(f"  [{j.get('score')}] {j.get('title','')[:45]} — {j.get('company','')[:24]}  -> {folder}/")

    save_done(args.seen_file, done)
    print(f"\nDone. Tailored {made} application(s) into {args.out}/.")


if __name__ == "__main__":
    main()
